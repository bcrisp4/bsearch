"""Data-driven docx/xlsx builders.

Office formats are zip archives whose entries carry timestamps, and both
python-docx and openpyxl stamp document properties with "now" — so every
output is normalised afterwards (fixed core properties, fixed zip entry
times) to keep regeneration byte-identical, matching the render pipeline's
determinism contract.
"""

from __future__ import annotations

import datetime
import io
import zipfile
from copy import copy
from typing import TYPE_CHECKING, Any

from docx import Document
from openpyxl import Workbook

if TYPE_CHECKING:
    from pathlib import Path

_EPOCH = datetime.datetime(2026, 1, 1, tzinfo=datetime.UTC)
_ZIP_TIME = (2026, 1, 1, 0, 0, 0)


def _normalise_zip(path: Path) -> None:
    """Rewrite a zip archive with fixed entry timestamps, sorted names."""
    with zipfile.ZipFile(path) as zin:
        entries = {name: zin.read(name) for name in zin.namelist()}
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in sorted(entries):
            info = zipfile.ZipInfo(name, date_time=_ZIP_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            zout.writestr(info, entries[name])
    path.write_bytes(buf.getvalue())


def build_docx_letter(entry: dict[str, Any], out_path: Path) -> None:
    """A simple business letter: sender block, date, recipient, body."""
    doc = Document()
    for line in entry["sender_lines"]:
        doc.add_paragraph(line)
    doc.add_paragraph("")
    doc.add_paragraph(entry["date"])
    doc.add_paragraph("")
    for line in entry["recipient_lines"]:
        doc.add_paragraph(line)
    doc.add_paragraph("")
    heading = doc.add_paragraph()
    heading.add_run(entry["subject"]).bold = True
    for para in entry["paragraphs"]:
        doc.add_paragraph(para)
    doc.add_paragraph("")
    for line in entry["signature_lines"]:
        doc.add_paragraph(line)

    props = doc.core_properties
    props.created = _EPOCH.replace(tzinfo=None)
    props.modified = _EPOCH.replace(tzinfo=None)
    props.last_modified_by = ""
    props.author = ""
    props.revision = 1
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    _normalise_zip(out_path)


def build_xlsx_sheet(entry: dict[str, Any], out_path: Path) -> None:
    """A single-sheet workbook of plain rows; first row bold."""
    wb = Workbook()
    ws = wb.active
    if ws is None:  # pragma: no cover — openpyxl contract
        raise RuntimeError("openpyxl workbook has no active sheet")
    ws.title = entry["sheet_name"]
    for row in entry["rows"]:
        ws.append(row)
    for cell in ws[1]:
        bold = copy(cell.font)
        bold.bold = True
        cell.font = bold  # pyright: ignore[reportAttributeAccessIssue] — openpyxl setter untyped

    wb.properties.created = _EPOCH.replace(tzinfo=None)
    wb.properties.modified = _EPOCH.replace(tzinfo=None)
    wb.properties.creator = ""
    wb.properties.lastModifiedBy = ""
    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))
    _normalise_zip(out_path)
